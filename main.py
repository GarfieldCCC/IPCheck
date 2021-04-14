import pymysql
from docx import Document
import win32com.client as wc
import os
import xlwt


class Method:
    def __init__(self):
        self.info = "success"

    def remove_brackets(self, list_):
        """
        去除括号

        :param list_: 带有括号的list
        :return: 没有括号的list
        :rtype: list
        """

        res = []
        for item in list_:
            res.append(item[0])
        self.info = "Successfully remove brackets! "
        return res

    def de_duplication(self, list_):
        """list去重

        :param list_: 去重前的list
        :return: 去重后的list
        :rtype: list
        """
        self.info = "Successfully de-duplication! "
        return list(set(list_))

    def compare(self, list_a, list_b):
        """
        list比较

        :param list_a: 文件读取的list
        :type list_a: list
        :param list_b: 数据库读取的list
        :type list_b: list
        :return: list_a中有但是list_b中没有的结果, list
        :rtype: list
        """
        self.info = "Successfully compare! "
        return list(set(list_a) - set(list_b))

    def console(self, info):
        self.info = "Successfully print! "
        print(info + " -- 文件有数据库没有: ")

    def output_to_excel(self, list_, list_em, path):
        """
        将比较结果输出至excel表格

        :param list_: 结果list
        :type list_: list
        :param list_em: 管理员名单
        :type list_em: list
        :param path: 输出路径
        :type path: str
        """
        workbook = xlwt.Workbook(encoding='utf-8')
        worksheet = workbook.add_sheet('结果')
        longest_ip = max(list_, key=len)
        longest_em = max(list_em, key=len)

        for i in range(list_.__len__()):
            worksheet.write(i + 1, 0, label=list_[i])
            worksheet.write(i + 1, 1, label=list_em[i])
        worksheet.col(0).width = 256 * (longest_ip.__len__() + 1)
        worksheet.col(1).width = 256 * (longest_em.__len__() + 1)
        workbook.save(path)
        self.info = "Successfully output! "

    def segmentation(self, ip):
        """
        网段切分: 前三段

        :param ip: ip地址
        :type ip: str

        :return: ip地址前三段: xxx.xxx.xxx.
        :rtype: str
        """
        ip_seg = ".".join(ip.split(".")[0:3]) + "."
        self.info = "Successfully seg! "
        return ip_seg

    def like_condition(self, ip):
        """
        生成like条件

        :param ip: ip地址
        :type ip: str

        :return: like 语句
        :rtype: str
        """
        return "like '" + self.segmentation(ip) + "%'"

    def equal_condition(self, employee):
        """
        生成like条件

        :param employee: ip地址
        :type employee: str

        :return: = 语句
        :rtype: str
        """
        self.info = "Successfully con! "
        return "= '" + employee + "'"


class SQL:
    def __init__(self):
        self.select = "select * from ip_record where area "
        self.select_ip = "select ip from ip_record where area "
        self.select_ip_seg = "select area, net_manager from ip_record where area "
        self.select_net_manager = "select username from t_user where employeeNum "
        self.method = Method()

    def search(self, area):
        """
        返回一条查询语句

        :param area: 区域
        :type area: str

        :return: "select * from ip_record where area = area"
        rtype: str
        """
        return self.select + "= '" + area + "'"

    def search_ip(self, area):
        """
        返回一条查询语句

        :param area: 区域
        :type area: str

        :return: "select ip from ip_record where area = area"
        rtype: str
        """
        return self.select_ip + "= '" + area + "'"

    def search_ip_segment(self, area, ip):
        """
        生成产业园ip网段查询的sql语句

        :param area: 区域
        :type area: str
        :param ip: ip地址
        :type ip: str

        :return: select area, net_manager from ip_record where area = 'area' and ip = xxx.xxx.xxx.%
        :rtype: str
        """
        sql_ = self.select_ip_seg + "= '" + area + "' and ip " + self.method.like_condition(ip)
        # print(sql_)
        return sql_

    def select_net_managers(self, employee):
        """
        生成代王ip网段查询的sql语句

        :param employee: 员工工号
        :type employee: str

        :return: 一段完整的sql语句
        :rtype: str
        """
        sql_ = self.select_net_manager + self.method.equal_condition(employee)
        # print(sql_)
        return sql_


class Docx:
    def __init__(self):
        self.length = 0

    def get_data(self, path):
        """读取文件

        :param path: 文件路径
        :type path: str

        :return: 一个ip的list
        :rtype: list
        """

        res = []
        document = Document(path)
        for p in document.paragraphs:
            if p.text.__len__() != 0:
                res.append(p.text.strip())
        self.length = res.__len__()
        return res

    def get_data_de_mac(self, path):
        """读取文件, 去掉Mac地址

        :param path: 文件路径
        :type path: str

        :return: 一个ip的list
        :rtype: list
        """

        flag = False
        # 如果是.doc文件
        if path[-1] == "c":
            # 获取当前路径
            c_path = os.getcwd()
            # 拼接为绝对路径
            a_path = c_path + "\\" + path.replace("/", "\\")

            word = wc.Dispatch("Word.Application")
            doc = word.Documents.Open(a_path)
            doc.SaveAs(a_path + "x", 12)
            word.Quit
            doc.close
            path = path + 'x'
            flag = True

        res = []
        document = Document(path)
        for p in document.paragraphs:
            if p.text.__len__() != 0:
                res.append(p.text.strip().split(" ")[0])
        self.length = res.__len__()

        if flag:
            os.remove(path)

        return res


class IPCheck:
    def __init__(self):
        self.host = "10.2.0.40"
        self.user = "root"
        self.password = "ipam@zll2020."
        self.database = "ipam"
        self.charset = "utf8"

    def execute_search(self, sql_select):
        """执行MySQL查询语句

        :param sql_select: sql查询语句
        :type sql_select: str

        :return: 返回查询结果的list
        :rtype: list
        """

        result = []
        conn = pymysql.connect(
            host=self.host,
            user=self.user,
            password=self.password,
            database=self.database,
            charset=self.charset
        )

        cursor = conn.cursor()

        cursor.execute(sql_select)

        res = cursor.fetchall()

        for r in res:
            result.append(r)

        cursor.close()

        conn.close()

        return result


def main():
    sql = SQL()
    docx = Docx()
    method = Method()
    ip_check = IPCheck()
    cyy = "产业园"
    jt = "集团"
    dw = "代王"
    output_path_cyy = "Excel/" + cyy + ".xls"
    output_path_jt = "Excel/" + jt + ".xls"
    output_path_dw = "Excel/" + dw + ".xls"
    path_ip_cyy = "Doc/ip-2021-4-12-cyy.doc"
    path_ip_jt = "Doc/ip-2021-4-12-jt.doc"
    path_ip_dw = "Doc/ip-2021-4-12-dw.docx"

    def process(file_path, output_path, area):
        """
        主流程

        :param file_path: 文件目录
        :type file_path: str
        :param output_path: 输出目录
        :type output_path: str
        :param area: 区域
        :type area: str
        """

        # 获取文件中的所有ip
        ip_docx_list = method.de_duplication(docx.get_data_de_mac(file_path))

        # 获取数据库中的所有ip
        ip_db_list = method.de_duplication(method.remove_brackets(ip_check.execute_search(sql.search_ip(area))))

        # 比较文件中的ip和数据库中的ip
        res = method.compare(ip_docx_list, ip_db_list)
        method.console(area)
        print(res.__len__(), res, "\n")
        res.sort()

        res_a = []
        for ip in res:
            sql_select = sql.search_ip_segment(area, ip)
            ip_seg_list = method.de_duplication(ip_check.execute_search(sql_select))
            employee_list = []
            if ip_seg_list.__len__() != 0:
                for em in ip_seg_list:
                    sql_select = sql.select_net_managers(em[1])
                    em_list = method.de_duplication(method.remove_brackets(ip_check.execute_search(sql_select)))
                    employee_list.append(em_list[0])
            res_a.append(", ".join(employee_list))

        # 将结果输出至excel表格
        method.output_to_excel(res, res_a, output_path)

    process(path_ip_cyy, output_path_cyy, cyy)
    process(path_ip_jt, output_path_jt, jt)
    process(path_ip_dw, output_path_dw, dw)


if __name__ == '__main__':
    test_sql = "select * from ip_record where area = '代王' and ip = '10.26.0.8'"

    main()
